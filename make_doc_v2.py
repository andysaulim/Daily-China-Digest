"""Generate CSIS_Digest_Presentation.docx — unified China + Korea + replication guide."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1B, 0x2A, 0x4A)
GOLD  = RGBColor(0xD4, 0xAC, 0x0D)
GRAY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MID   = RGBColor(0x88, 0x88, 0x88)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def para(doc, text, bold=False, italic=False, size=11, color=None,
         align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def heading(doc, text, level=1):
    sizes = {1: 17, 2: 13, 3: 11}
    p = para(doc, text, bold=True, size=sizes.get(level, 11), color=NAVY,
             space_before=16 if level == 1 else 10, space_after=4)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1B2A4A')
        pBdr.append(bot); pPr.append(pBdr)
    return p


def bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + "  ")
        r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY


def callout(doc, label, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    pBdr.append(left); pPr.append(pBdr)
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '1B2A4A')
        p = cell.paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(9); run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, 'F4F6F9')
            p = cell.paragraphs[0]
            bold = isinstance(val, tuple)
            text_val = val[0] if bold else val
            run = p.add_run(text_val)
            run.bold = bold and val[1]
            run.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ── BUILD ────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── COVER ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("CSIS Korea Chair")
r.font.size = Pt(10); r.font.color.rgb = GOLD; r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Automated Daily Intelligence Digest")
r.font.size = Pt(26); r.font.color.rgb = NAVY; r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("China Brief  ·  Korea Brief")
r.font.size = Pt(15); r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
r = p.add_run("What It Is · How It Works · Cost · Replication · Q&A")
r.font.size = Pt(11); r.font.color.rgb = MID; r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Andy Lim  ·  CSIS Korea Chair  ·  May 2026")
r.font.size = Pt(10); r.font.color.rgb = MID

doc.add_page_break()

# ── 1. WHAT IS THIS ──────────────────────────────────────────────────────────
heading(doc, "1.  What Is This?")
para(doc,
     "Two automated daily intelligence briefings — one on China, one on Korea — that deliver "
     "a Presidential Daily Brief-style email to your inbox at 6:30 AM ET every morning. "
     "Each covers overnight news, expert analysis, live market data, government actions, and "
     "satellite imagery flags across 130–138 sources. It takes 5 minutes to read and replaces "
     "the 90 minutes a morning it would take an analyst to do the same manually.",
     size=11, color=GRAY, space_after=8)

callout(doc, "The one-line pitch:",
        "A senior analyst briefing, automated. No staff time. No manual curation. "
        "Runs entirely on its own, every day, for about $12–15/month per digest.")

table(doc,
    ["", "China Daily Brief", "Korea Daily Brief"],
    [
        ("Sources", "130+", "138+"),
        ("Tiers", "4 (News / Analysis / Academic / PRC State Media)", "4 (News / Analysis / Academic / KCNA)"),
        ("Languages", "English + Chinese (translated by Claude)", "English + Korean (translated by Claude)"),
        ("Sections", "14 active sections", "9 sections + trackers"),
        ("Market data", "SSE, Hang Seng, USD/CNY, CGB, Brent, PBOC LPR", "KOSPI, KRW/USD, BOK rate, 10Y yield"),
        ("Persistent trackers", "Xi Jinping appearances", "Kim Jong Un appearances + 11 facility statuses"),
        ("Delivery", "6:30 AM ET daily", "6:00 AM ET daily"),
        ("Monthly API cost", "~$5–8", "~$5–8"),
    ],
    col_widths=[1.4, 2.6, 2.6]
)

# ── 2. HOW IT WORKS ──────────────────────────────────────────────────────────
heading(doc, "2.  How It Works")
para(doc, "Every morning, GitHub's free servers wake up and run a 4-step pipeline in 3–5 minutes:", size=11, color=GRAY, space_after=6)

bullet(doc, "Scans 130–138 RSS feeds simultaneously — wire services, think tanks, academic journals, state media, government sites, Korean/Chinese-language press", bold_prefix="Step 1 — Collect")
bullet(doc, "Sends all articles to Claude (Anthropic's AI) with a ~20,000-token prompt. Claude reads, translates non-English sources, and writes the full structured digest as JSON", bold_prefix="Step 2 — Generate")
bullet(doc, "A validation gate checks the output before anything sends: word count floor, section completeness, URL validity, duplicate detection, hallucination guards. Retries up to 2× if it fails", bold_prefix="Step 3 — Validate")
bullet(doc, "Converts the JSON to a styled HTML email and sends via Gmail. Archives the issue to GitHub Pages for a public record", bold_prefix="Step 4 — Send")

para(doc, "", space_after=4)
callout(doc, "Fallback logic:",
        "If the 6:30 AM run fails, a 7:30 AM backup fires automatically. "
        "If both fail, no email sends — silence is better than a bad product.")

table(doc,
    ["Tool", "Role", "Cost"],
    [
        ("Python", "Pipeline code", "Free"),
        ("GitHub Actions", "Runs on a timer — no server needed", "Free"),
        ("Claude API (Anthropic)", "Reads, translates, and writes the digest", "~$5–8/month per digest"),
        ("Gmail SMTP", "Sends the email", "Free"),
        ("GitHub Pages", "Public archive of past issues", "Free"),
        ("Claude Code", "AI coding assistant used to build the pipeline", "~$20/month (dev only)"),
    ],
    col_widths=[1.6, 3.2, 1.5]
)

# ── 3. WHAT IT COVERS ────────────────────────────────────────────────────────
heading(doc, "3.  What Each Issue Covers")

para(doc, "China Daily Brief", bold=True, size=12, color=NAVY, space_after=4)
table(doc,
    ["Section", "Content"],
    [
        ("Market Strip", "SSE Composite, Hang Seng, USD/CNY, Brent, 10Y CGB, PBOC rates — with daily moves"),
        ("Δ Since Yesterday", "What changed overnight: tariff actions, PLA flights, coast guard presence"),
        ("Morning Memo", "Top 3 stories in 2 sentences — the elevator brief"),
        ("Top Stories", "2–4 major items with 'So what' analysis and pattern note"),
        ("Overnight Flash", "6 secondary overnight items"),
        ("Expert Analysts", "What CSIS, CFR, Brookings, Sinocism, Foreign Affairs published today — linked"),
        ("Satellite & Location Watch", "AMTI gray-zone activity flags + Hidden Reach overseas footprint"),
        ("PRC Government", "State Council, MOFA, MOFCOM, personnel changes"),
        ("US–China Trade & Sanctions", "Section 301/232/IEEPA tariff stack, Entity List adds, CFIUS"),
        ("Congressional Watch", "Select Committee on CCP, SFRC, USCC, CECC"),
        ("Business & Economy", "Major corporates, macro, property, tech"),
        ("Indo-Pacific", "Cross-Strait, Japan, Philippines, Australia, India, Vietnam"),
        ("Also Today / The Wire", "Everything else worth flagging"),
        ("On This Day", "A verified historical China event matching today's date"),
    ],
    col_widths=[2.0, 4.2]
)

para(doc, "Korea Daily Brief", bold=True, size=12, color=NAVY, space_after=4, space_before=8)
table(doc,
    ["Section", "Content"],
    [
        ("Market Strip", "KOSPI, KRW/USD, BOK rate, 10Y yield — with daily moves"),
        ("Morning Memo", "Top 3 overnight items at a glance"),
        ("Top Stories", "2–4 major items with 'So what' analysis"),
        ("KCNA Delta", "North Korean state media output — watch flags, silence signals, key statements"),
        ("Kim Jong Un Tracker", "Appearance status with 7-day anomaly flag if absent"),
        ("Facility Watch", "11 monitored sites: Yongbyon, Punggye-ri, Sohae, Sinpo, border crossings"),
        ("Expert Analysts", "38 North, CSIS, Carnegie, Stimson, War on the Rocks — linked articles"),
        ("US-ROK-Japan Tracker", "Alliance developments, USFK, trilateral coordination"),
        ("The Wire", "Secondary items and also-today summaries"),
    ],
    col_widths=[2.0, 4.2]
)

# ── 4. COSTS ─────────────────────────────────────────────────────────────────
heading(doc, "4.  Costs")
para(doc,
     "The only paid dependency is the Anthropic API. Everything else — GitHub, Gmail, the archive — is free.",
     size=11, color=GRAY, space_after=8)

table(doc,
    ["Item", "Per Run", "Per Month"],
    [
        ("Claude Sonnet (normal day)", "~$0.13", "~$4 (if every run succeeds first try)"),
        ("Claude Opus (retry day)", "~$0.65", "Kicks in ~20% of runs"),
        ("Realistic blended cost", "~$0.20", "~$5–8 per digest"),
        ("Both digests combined", "", "~$10–15 API"),
        ("GitHub Actions + Pages + Gmail", "", "$0"),
        ("Claude Code (development only)", "", "~$20 (cancel when not making changes)"),
        (("TOTAL — both digests running", True), "", ("~$30–35/month", False)),
    ],
    col_widths=[2.5, 1.5, 2.7]
)

callout(doc, "How the model logic works:",
        "The pipeline tries the cheaper Sonnet model first (~$0.13/run). "
        "If the output doesn't meet quality minimums — word count, section completeness, source diversity — "
        "it automatically retries with Opus (~$0.65/run). Most days Sonnet is sufficient.")

# ── 5. HALLUCINATION ─────────────────────────────────────────────────────────
heading(doc, "5.  Hallucination — The Credibility Problem and How We Solve It")
para(doc,
     "Hallucination is when AI generates plausible-sounding but false content — invented quotes, "
     "fabricated article titles, links that go nowhere. For a briefing read by NSC staff, "
     "Korea desk officers, and senior scholars, one wrong fact destroys credibility. "
     "Both digests run 9 guards:",
     size=11, color=GRAY, space_after=8)

bullet(doc, "Claude is only allowed to use articles it was actually given as input — it cannot invent sources", bold_prefix="1. Grounded inputs")
bullet(doc, "Every URL must come verbatim from the source data — Claude is explicitly told never to construct or modify a URL", bold_prefix="2. URL discipline")
bullet(doc, "Google News redirect links are resolved to real article URLs before Claude sees them — unresolvable links are stripped before sending", bold_prefix="3. Link pre-resolution")
bullet(doc, "Word count floor, required sections present, no placeholder text, date matches today — checked before every send", bold_prefix="4. Pre-send validation gate")
bullet(doc, "Think tank and journal citations must come from the actual RSS feeds — Claude cannot add expert analysis that wasn't in the input", bold_prefix="5. Expert analyst grounding")
bullet(doc, "News articles are blocked from appearing in the academic journal section via domain blocklist + URL checks", bold_prefix="6. Academic tier filter")
bullet(doc, "~35 news outlet domains blocked from the academic tier — prevents headlines bleeding into journal section via keyword match", bold_prefix="7. News domain blocklist")
bullet(doc, "If the validation gate fails, the pipeline retries digest generation up to 2× with the failure reason injected into the prompt", bold_prefix="8. Auto-retry with feedback")
bullet(doc, "No classified sources, no paywalled content, no signals intelligence — only what is publicly and freely available", bold_prefix="9. Scope limits")

para(doc, "", space_after=4)
callout(doc, "Residual risk:",
        "The 'So what' analysis and pattern notes are AI interpretations, not sourced facts. "
        "Headlines, summaries, and market data are all grounded in real sources. "
        "Treat the analysis as a prompt for thinking, not a finished judgment.",
        color=RED)

# ── 6. REPLICATION ───────────────────────────────────────────────────────────
heading(doc, "6.  Replication — Building a Third Digest")
para(doc,
     "The pipeline architecture is topic- and region-agnostic. "
     "The hard infrastructure — scheduler, email delivery, validation gate, HTML rendering, "
     "market strip, hallucination guards — stays exactly the same. "
     "Only the content layer changes.",
     size=11, color=GRAY, space_after=8)

table(doc,
    ["File", "What You Change", "Time"],
    [
        ("collect.py", "Swap in new RSS feeds for the target region/topic", "1–2 days"),
        ("digest.py", "Rewrite the Claude prompt with the right analytical frame and audience", "Half a day"),
        ("render.py", "Update branding, section names, header", "2–3 hours"),
        ("send_email.py", "Update recipient list", "10 minutes"),
    ],
    col_widths=[1.2, 3.5, 1.0]
)

para(doc, "Example: An Africa Digest", bold=True, size=12, color=NAVY, space_after=6, space_before=8)
para(doc,
     "Africa is an ideal next candidate — strong English-language press, growing China/Russia "
     "engagement that CSIS already tracks, and significant gaps in existing automated monitoring.",
     size=11, color=GRAY, space_after=6)

table(doc,
    ["Tier", "Africa Sources", "Window"],
    [
        ("Tier 1 — News", "Reuters Africa, AP, BBC Africa, Al Jazeera Africa, VOA Africa, AllAfrica, Daily Maverick, Premium Times (Nigeria), The East African, Daily Nation (Kenya), Mail & Guardian", "24h"),
        ("Tier 2 — Analysis", "Africa Center for Strategic Studies, Brookings Africa, SAIS Africa, ISS Africa, Chatham House Africa, Council on Foreign Relations Africa, War on the Rocks, Foreign Policy", "36h"),
        ("Tier 3 — Academic", "African Affairs, Journal of Modern African Studies, African Security, Journal of Eastern African Studies", "72h"),
        ("Tier 4 — State Media", "CCTV Africa, Xinhua Africa, RT Africa, China Daily Africa, Sputnik Africa", "24h"),
    ],
    col_widths=[1.3, 3.9, 0.7]
)

para(doc, "Africa-specific sections would include:", size=11, color=GRAY, space_after=4, space_before=4)
bullet(doc, "Sahel security tracker — military activity, coup watch, Wagner/Russia presence in Mali, Niger, Burkina Faso", bold_prefix="Security:")
bullet(doc, "China-Africa engagement — BRI lending, port deals, Huawei contracts, diplomatic activity", bold_prefix="China footprint:")
bullet(doc, "Critical minerals — cobalt (DRC), copper (Zambia), lithium, rare earths — supply chain flags", bold_prefix="Resources:")
bullet(doc, "AU / ECOWAS / IGAD — regional body activity, peacekeeping mandates, election watch", bold_prefix="Institutions:")
bullet(doc, "Commodity prices — copper, cobalt, gold, crude — tied to market strip", bold_prefix="Markets:")

para(doc, "", space_after=4)
callout(doc, "Realistic estimate:",
        "A working Africa digest could be operational in 3–5 days of part-time work, "
        "built directly on top of the existing China or Korea codebase. "
        "Monthly cost would be identical: ~$5–8/month in API fees.")

# ── 7. DO YOU NEED CLAUDE CODE ───────────────────────────────────────────────
heading(doc, "7.  Do You Need Claude Code After It's Running?")
para(doc,
     "No. Claude Code is a development tool — it was used to build, debug, and improve the pipeline. "
     "Once it's set up, it runs fully automatically. Claude Code is not part of the daily operation.",
     size=11, color=GRAY, space_after=8)

table(doc,
    ["Scenario", "Need Claude Code?"],
    [
        ("Daily operation — pipeline runs, email delivers", "No"),
        ("Adding a new RSS source", "No — edit one line in collect.py"),
        ("Changing the recipient list", "No — update one environment variable"),
        ("Adjusting the send time", "No — change one line in the schedule file"),
        ("Adding a new digest section", "Yes — prompt engineering + render changes"),
        ("Building a third digest (new region)", "Yes — 3–5 days"),
        ("Debugging a pipeline failure", "Helpful, not required"),
    ],
    col_widths=[3.8, 2.4]
)

callout(doc, "Ongoing cost without Claude Code:",
        "Cancel the ~$20/month Claude Code subscription once development is complete. "
        "The pipeline continues running on just the Claude API (~$10–15/month for both digests).")

# ── 8. Q&A ───────────────────────────────────────────────────────────────────
heading(doc, "8.  Anticipated Questions")

qa = [
    ("What if the AI gets something wrong?",
     "The validation gate catches the most common failures before the email sends. "
     "For analytical errors in the 'So what' framing, the pipeline doesn't self-correct — "
     "a human reader would flag it. This is why the interpretive analysis should be treated "
     "as a prompt, not a verdict. The factual layer (headlines, sources, URLs, market data) "
     "is grounded and checked."),

    ("Could this replace a research assistant?",
     "For the read-ahead function — scanning 130+ sources overnight and producing a structured brief — yes, substantially. "
     "For deeper analytical work, stakeholder outreach, or judgment calls, no. "
     "The Korea and China digests replace the first 90 minutes of an analyst's morning, "
     "not the analyst."),

    ("Why Claude and not GPT-4 or Gemini?",
     "Claude was chosen for: (1) large context window that handles 130+ articles per run, "
     "(2) reliable structured JSON output, (3) prompt caching that reduces repeat costs, "
     "(4) strong performance on Korean and Chinese-language content. "
     "The system is model-agnostic in principle — changing the model is a one-line edit."),

    ("How does Korean/Chinese-language content work?",
     "About 20 of the Korea digest's 71 Tier 1 feeds are Korean-language sources "
     "(조선일보, 한겨레, JTBC, KBS, etc.). The China digest includes 人民日报, 新华社, 环球时报, and Caixin. "
     "These articles are collected in their original language and sent to Claude, "
     "which translates and analyzes them during digest generation. "
     "This gives coverage that English-only monitoring misses entirely."),

    ("Is this secure? Are our sources exposed?",
     "The pipeline reads only publicly available RSS feeds and open websites. "
     "No credentials, subscriptions, or classified systems are accessed. "
     "The Anthropic API key and Gmail password are stored as encrypted GitHub secrets — "
     "not in the code itself."),

    ("What if Anthropic raises prices or Claude gets worse?",
     "The model name is a single variable in the code. Switching to a different Claude model, "
     "or in principle a different AI provider, requires changing one line. "
     "The pipeline is not locked to any specific model version."),

    ("What happens if the pipeline fails?",
     "A fallback cron fires 60 minutes after the primary run. "
     "If the primary already succeeded, the fallback detects this and skips. "
     "If both fail, no email sends — silence is preferable to a broken product. "
     "GitHub Actions logs show exactly where and why it failed."),

    ("How long did this take to build?",
     "The Korea digest took about two weeks of part-time development to reach a reliable daily cadence. "
     "The China digest was built in roughly one week on top of the Korea codebase. "
     "An Africa digest would take 3–5 days — the infrastructure is already proven."),

    ("Can non-technical staff manage this once it's running?",
     "Yes. Day-to-day operation requires no technical knowledge — it runs itself. "
     "Adding recipients, adjusting timing, or enabling/disabling the workflow "
     "can all be done through GitHub's web interface with no coding. "
     "Code changes require a developer or Claude Code."),

    ("Could this be distributed to a larger list?",
     "Yes. The recipient list is a comma-separated environment variable. "
     "It can send to any number of addresses — a list of 50 costs the same as a list of 1. "
     "Future work includes subscriber management with subscribe/unsubscribe."),

    ("What's the biggest ongoing risk?",
     "Source quality. The digest is only as good as its inputs. "
     "If a key source goes behind a paywall, changes its RSS format, or stops publishing, "
     "that coverage disappears without notification. "
     "Quarterly review of source coverage is the main ongoing maintenance task."),
]

for q, a in qa:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q:  " + q)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.15)
    r2 = p2.add_run(a)
    r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY

# ── BACK COVER ───────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run("CSIS Korea Chair  ·  Daily Intelligence Digest")
r.font.size = Pt(12); r.font.color.rgb = NAVY; r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("China Brief  ·  Korea Brief  ·  Replication Guide")
r.font.size = Pt(11); r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Code: github.com/andysaulim/Daily-China-Digest")
r.font.size = Pt(10); r.font.color.rgb = MID

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Andy Lim  ·  May 2026")
r.font.size = Pt(10); r.font.color.rgb = MID

doc.save("CSIS_Digest_Presentation.docx")
print("Saved CSIS_Digest_Presentation.docx")
