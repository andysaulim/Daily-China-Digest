"""Generate CSIS_Digest_Overview.docx presentation document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xD4, 0xAC, 0x0D)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED  = RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xF5, 0xF5, 0xF5)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def para(doc, text, bold=False, size=11, color=None, align=None, space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = para(doc, text, bold=True, size=sizes.get(level, 12), color=NAVY,
             space_before=14 if level == 1 else 10, space_after=4)
    # Add bottom border for H1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1B2A4A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1B2A4A')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, 'F4F6F9')
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                text, bold = val
            else:
                text, bold = val, False
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def callout(doc, label, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    pBdr.append(left)
    pPr.append(pBdr)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    return p

# ── Build document ──────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── COVER ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
r = p.add_run("CSIS Korea Chair")
r.font.size = Pt(10)
r.font.color.rgb = GOLD
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Daily Intelligence Digest")
r.font.size = Pt(26)
r.font.color.rgb = NAVY
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("China Brief  ·  Korea Brief")
r.font.size = Pt(14)
r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(32)
r = p.add_run("How It Works — Setup, Replication, and Costs")
r.font.size = Pt(11)
r.font.color.rgb = GRAY
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared by Andy Lim  ·  May 2026")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ── 1. WHAT IS THIS ─────────────────────────────────────────────────────────
heading(doc, "1.  What Is This?")
para(doc,
     "Two automated daily intelligence briefings — one on China, one on Korea — that land in your inbox "
     "at 6:30 AM ET every morning. Each covers overnight news, expert analysis, market indicators, "
     "government actions, and satellite imagery flags across 130+ sources. Formatted like a professional "
     "read-ahead. No manual curation. No staff time. Runs entirely on its own.",
     size=11, color=GRAY, space_after=8)

callout(doc, "Bottom line:", "A senior analyst briefing, automated. Each issue takes about 5 minutes to read and covers everything a China or Korea watcher needs before their first meeting.")

# ── 2. HOW IT WORKS ─────────────────────────────────────────────────────────
heading(doc, "2.  How It Works — The Pipeline")
para(doc, "Every morning at 6:30 AM, GitHub's free servers wake up and run four steps:", size=11, color=GRAY, space_after=6)

bullet(doc, "Scans 130+ RSS feeds simultaneously — wire services, think tanks, academic journals, PRC state media, government sites", bold_prefix="Step 1 — Collect")
bullet(doc, "Sends all the articles to Claude (Anthropic's AI) with a detailed set of instructions. Claude reads everything and writes the digest as structured data", bold_prefix="Step 2 — Generate")
bullet(doc, "Converts Claude's output into a styled HTML email — market strip, section labels, linked headlines, analysis cards", bold_prefix="Step 3 — Render")
bullet(doc, "Emails the finished digest via Gmail to the distribution list", bold_prefix="Step 4 — Send")

para(doc, "", space_after=4)
para(doc, "The whole process takes about 3–5 minutes and costs roughly $0.13–$0.65 per issue depending on how much AI processing is needed that day.", size=11, color=GRAY, space_after=10)

add_table(doc,
    ["Component", "What It Does", "Cost"],
    [
        ("Python", "The programming language the pipeline is written in", "Free"),
        ("GitHub Actions", "Runs the pipeline on a timer — no server or IT infrastructure needed", "Free"),
        ("Claude API (Anthropic)", "AI that reads all articles and writes the digest", "~$5–8/month per digest"),
        ("Gmail SMTP", "Sends the finished email", "Free"),
        ("GitHub Pages", "Hosts a public archive of past issues", "Free"),
        ("Claude Code", "AI coding assistant used to build and modify the pipeline", "~$20/month (development only)"),
    ],
    col_widths=[1.5, 3.2, 1.5]
)

# ── 3. WHAT IT COVERS ───────────────────────────────────────────────────────
heading(doc, "3.  What Each Digest Covers")

para(doc, "China Daily Brief", bold=True, size=12, color=NAVY, space_after=4)
add_table(doc,
    ["Section", "Content"],
    [
        ("Market Strip", "SSE Composite, Hang Seng, USD/CNY, Brent crude, 10-year CGB, PBOC rates — with daily moves"),
        ("Δ Since Yesterday", "What changed: tariff actions, PLA flights, coast guard presence, rate moves"),
        ("Morning Memo", "Top 3 stories in 2 sentences each — the elevator brief"),
        ("Top Stories", "2–4 major hard news items with 'So what' and pattern analysis"),
        ("Overnight Flash", "6 secondary items from overnight"),
        ("Expert Analysts", "What CSIS, CFR, Brookings, Sinocism, Foreign Affairs published today — with links"),
        ("Satellite & Location Watch", "Activity flags from AMTI gray-zone sites and Hidden Reach overseas footprint"),
        ("PRC Government", "State Council, MOFA, MOFCOM, personnel changes"),
        ("US–China Trade & Sanctions", "Section 301/232/IEEPA tariff stack, Entity List adds, CFIUS actions"),
        ("Congressional Watch", "Select Committee on CCP, SFRC, USCC, CECC activity"),
        ("Business & Economy", "Major corporates, macro indicators, property, tech sector"),
        ("Indo-Pacific", "Cross-Strait, Japan, Philippines, Australia, India, Korea, Vietnam"),
        ("Also Today / The Wire", "Everything else worth noting — short summaries"),
        ("On This Day", "A verified historical China event matching today's exact date"),
    ],
    col_widths=[2.0, 4.2]
)

para(doc, "Korea Daily Brief covers:", bold=True, size=11, color=NAVY, space_after=4)
para(doc, "Korean Peninsula security (North Korea, US-ROK alliance, USFK), ROK domestic politics and economy, "
     "KCNA / Rodong Sinmun state media delta, Kim Jong Un appearance tracking with anomaly flag, "
     "and the US-ROK-Japan trilateral tracker. Identical pipeline architecture.", size=11, color=GRAY, space_after=10)

# ── 4. SOURCES ──────────────────────────────────────────────────────────────
heading(doc, "4.  Source Coverage")
add_table(doc,
    ["Tier", "Count", "Examples", "Lookback"],
    [
        ("News", "85+ feeds", "Reuters, AP, Bloomberg, NYT, WSJ, FT, Nikkei, SCMP, Xinhua, Global Times", "24 hours"),
        ("Analysis", "28 feeds", "CSIS, CFR, Brookings, Carnegie, MERICS, RAND, Sinocism, Stimson, ASPI", "36 hours"),
        ("Academic", "16 feeds", "International Security, China Quarterly, Journal of Contemporary China", "72 hours"),
        ("PRC Primary", "8 feeds", "人民日报, 新华社, 环球时报, CCTV, MOFA Presser, State Council, TAO", "24 hours"),
    ],
    col_widths=[1.0, 0.7, 3.5, 1.0]
)
callout(doc, "Prestige rule:", "Stories from WSJ, NYT, WaPo, Bloomberg, FT, The Economist, Reuters, CNN, CNBC are never dropped regardless of AI prioritization. They are guaranteed inclusion.")

# ── 5. COSTS ────────────────────────────────────────────────────────────────
heading(doc, "5.  Costs")
add_table(doc,
    ["Item", "Monthly Cost", "Notes"],
    [
        ("Claude API — China digest", "~$5–8", "~$0.13/run on Sonnet; ~$0.65/run if Opus retry needed"),
        ("Claude API — Korea digest", "~$5–8", "Same model logic, slightly smaller payload"),
        ("GitHub Actions + Pages", "$0", "Free for public repositories"),
        ("Gmail SMTP", "$0", "Free with a standard Gmail account"),
        ("Claude Code (dev tool)", "~$20", "Only needed if making changes to the pipeline"),
        ("TOTAL (both digests running)", "~$30–35/month", ""),
    ],
    col_widths=[2.2, 1.3, 2.7]
)
callout(doc, "How the AI pricing works:",
        "The pipeline tries the cheaper model (Sonnet, ~$0.13/run) first. "
        "If the output doesn't meet quality minimums — word count, section completeness, source diversity — "
        "it automatically retries with the more powerful model (Opus, ~$0.65/run). "
        "Most days Sonnet is sufficient.")

# ── 6. HALLUCINATION ────────────────────────────────────────────────────────
heading(doc, "6.  Hallucination — What It Is and How We Control It")
para(doc,
     "Hallucination is when an AI model generates plausible-sounding but false information — "
     "fabricated names, invented quotes, made-up article titles, or links that go nowhere. "
     "This is the primary credibility risk for any AI-generated briefing.",
     size=11, color=GRAY, space_after=8)

para(doc, "What we do about it:", bold=True, size=11, color=NAVY, space_after=4)
bullet(doc, "Claude is only allowed to use articles it was actually given as input. It cannot invent sources.", bold_prefix="Grounded inputs:")
bullet(doc, "Every URL in the digest must come verbatim from the source articles. Claude is explicitly instructed never to construct or modify a URL.", bold_prefix="URL discipline:")
bullet(doc, "All Google News redirect links are resolved to the real article URL before Claude sees them. Any unresolvable links are stripped before the email sends.", bold_prefix="Link verification:")
bullet(doc, "The digest is checked before sending: minimum word count, required sections present, no 'None' strings, no placeholder text, date matches today.", bold_prefix="Pre-send validation:")
bullet(doc, "Think tank citations (CSIS, CFR, etc.) must come from the actual RSS feeds — Claude cannot add expert analysis entries that weren't in the input.", bold_prefix="Expert analyst grounding:")
bullet(doc, "No signals intelligence, no paywalled content, no classified sources. The pipeline only reads what is publicly and freely available.", bold_prefix="Scope limits:")

para(doc, "", space_after=4)
callout(doc, "Residual risk:",
        "The 'So what' analysis and pattern notes are AI-generated interpretations, not sourced facts. "
        "They are intended as prompts for thinking, not finished analytical judgments. "
        "Headlines and article summaries are grounded in real sources.",
        color=RED)

# ── 7. REPLICATION ──────────────────────────────────────────────────────────
heading(doc, "7.  Replicating for Another Country or Topic")
para(doc,
     "The pipeline is country-agnostic. The hard infrastructure — scheduler, email delivery, "
     "validation, HTML rendering, market strip — stays exactly the same. "
     "What changes is the content layer.",
     size=11, color=GRAY, space_after=8)

add_table(doc,
    ["File", "What You Change", "Time"],
    [
        ("collect.py", "Swap in new RSS feeds for the target country/topic", "1–2 days"),
        ("digest.py", "Rewrite Claude's instructions with the right analytical frame (e.g. India-Pakistan dynamics, Gulf security)", "Half a day"),
        ("render.py", "Update branding, header name, section labels", "A few hours"),
        ("send_email.py", "Update recipient list", "10 minutes"),
    ],
    col_widths=[1.3, 3.5, 1.0]
)

para(doc, "What would be harder:", bold=True, size=11, color=NAVY, space_after=4)
bullet(doc, "WSJ, FT, Bloomberg full-text requires subscriptions. We use RSS headlines and summaries (publicly available for free).", bold_prefix="Paywalled sources:")
bullet(doc, "Claude reads Chinese, Korean, Arabic, Russian, Farsi natively, but good public RSS feeds need to exist for the target language.", bold_prefix="Non-English primary sources:")
bullet(doc, "The more niche the subject, the fewer RSS feeds exist. Very specialized topics (specific weapons programs, subnational politics) may need manual source curation.", bold_prefix="Highly specialized topics:")

para(doc, "", space_after=4)
callout(doc, "Realistic estimate:", "A Japan digest or Southeast Asia digest could be running in 3–5 days of part-time work, built on top of the existing codebase.")

# ── 8. DO YOU NEED CLAUDE CODE AFTER SETUP ──────────────────────────────────
heading(doc, "8.  Do You Need Claude Code After Setup?")
para(doc,
     "No. Once the pipeline is set up and running, it runs fully automatically. "
     "Claude Code is a development tool — it was used to build, debug, and improve the pipeline. "
     "It is not part of the daily operation.",
     size=11, color=GRAY, space_after=8)

add_table(doc,
    ["Scenario", "Need Claude Code?"],
    [
        ("Daily operation — pipeline runs, email sends", "No"),
        ("Adding a new RSS source feed", "No — edit one line in collect.py"),
        ("Changing the recipient email list", "No — update an environment variable"),
        ("Adjusting the send time", "No — change one line in the GitHub Actions schedule"),
        ("Adding a new digest section", "Yes — requires prompt engineering and render changes"),
        ("Debugging a pipeline failure", "Helpful but not required"),
        ("Building a third digest (new country)", "Yes — 3–5 days of development"),
    ],
    col_widths=[3.8, 1.5]
)

callout(doc, "Ongoing cost without Claude Code:",
        "If no changes are being made, the ~$20/month Claude Code subscription can be cancelled. "
        "The pipeline continues running on just the Claude API (~$10–15/month for both digests).")

# ── 9. ANTICIPATED QUESTIONS ─────────────────────────────────────────────────
heading(doc, "9.  Anticipated Questions")

qa_pairs = [
    ("What happens if the AI gets something wrong?",
     "The validation gate catches the most common failures before the email sends — "
     "missing sections, placeholder text, broken links. For analytical errors (e.g. a mischaracterized 'So what'), "
     "the pipeline doesn't self-correct; a human reader would need to flag it. "
     "This is why the 'So what' analysis should be treated as a prompt, not a verdict."),
    ("Could this be used for internal research distribution?",
     "Yes — the recipient list is just a comma-separated environment variable. "
     "It can send to any number of email addresses. A distribution list of 20 people costs the same as a list of 1."),
    ("Is this secure / are our sources exposed?",
     "The pipeline reads only publicly available RSS feeds and open websites. "
     "No credentials, subscriptions, or classified systems are accessed. "
     "The Anthropic API key and Gmail password are stored as encrypted GitHub secrets, "
     "not in the code."),
    ("What if Anthropic changes their pricing or Claude gets worse?",
     "The pipeline is model-agnostic — the model name is a single variable. "
     "Switching to a different Claude model, or in principle a different AI provider, "
     "requires changing one line of code."),
    ("Could this replace a research assistant?",
     "For the read-ahead function — scanning overnight news and producing a structured brief — yes, substantially. "
     "For deeper analytical work, follow-up research, stakeholder outreach, or judgment calls, no. "
     "It is best understood as freeing an analyst's morning for higher-value work."),
    ("How long did it take to build?",
     "The Korea digest took about two weeks of part-time development to reach a reliable daily cadence. "
     "The China digest was built in roughly one week on top of the Korea codebase. "
     "Iteration has been ongoing — adding sections, improving source coverage, refining the prompt."),
    ("What does the email actually look like?",
     "A styled HTML email roughly 1,200–1,400 words: dark navy header, market strip, "
     "numbered morning memo, white story cards with left-border accents, expert analyst cards with linked headlines, "
     "tracker sections for government/trade/satellite. Renders correctly in Gmail, Outlook, and Apple Mail."),
    ("Could this be done without coding knowledge?",
     "Building it from scratch requires Python and some familiarity with APIs. "
     "Replicating an existing digest for a new country — swapping feeds and rewriting the prompt — "
     "is more accessible, especially with Claude Code guiding the edits. "
     "A motivated non-programmer could get there with a few days of guided work."),
    ("What is the biggest risk of this approach?",
     "Source quality. The digest is only as good as its inputs. "
     "If a key source goes behind a paywall, changes its RSS format, or stops publishing, "
     "that coverage disappears without notification. "
     "Regular review of source coverage is the main ongoing maintenance task."),
]

for q, a in qa_pairs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q: " + q)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.15)
    r2 = p2.add_run(a)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY

# ── FOOTER ──────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("CSIS Korea Chair  ·  Daily Intelligence Digest  ·  May 2026")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Code: github.com/andysaulim/Daily-China-Digest  ·  Prepared by Andy Lim")
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

doc.save("CSIS_Digest_Overview.docx")
print("Saved CSIS_Digest_Overview.docx")
